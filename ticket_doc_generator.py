import os
from datetime import datetime
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

exe_dir_path = os.path.dirname(os.path.abspath(__file__))


class TicketDocGenerator:
    def __init__(self, tickets: int, no_of_questions: int, questions: Iterable):
        self.tickets = tickets
        self.no_of_questions = no_of_questions
        self.questions = questions

    def generate(self):
        document = Document()
        current_date = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")
        for ticket in range(1, self.tickets + 1):
            title_paragraph = document.add_paragraph()
            title_paragraph.add_run(f"Билет № {ticket}").bold = True
            font = title_paragraph.style.font
            font.size = Pt(14)
            font.name = "Times New Roman"
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            list_paragraph = document.add_paragraph()
            list_paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
            for i, question in enumerate(self.questions[ticket - 1], start=1):
                list_paragraph.add_run(f"{i}.\t{question}")
                list_paragraph.add_run("\n")
            if self.tickets > 1 and ticket != self.tickets + 1:
                document.add_page_break()

            document.save(f"Testes_{current_date}.docx")
